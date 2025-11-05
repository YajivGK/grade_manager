import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except Exception as e:
    raise SystemExit("Missing dependency: python-docx. Install with `pip install python-docx`.")


def set_run_font(run, size_pt=12, bold=False):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size_pt)
    run.bold = bold


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    size = 14 if level == 1 else 12
    set_run_font(r, size_pt=size, bold=True)
    return p


def add_body_paragraph(doc, text=""):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size_pt=12)
    # 1.5 line spacing
    p.paragraph_format.line_spacing = 1.5
    return p


def add_code_block(doc, code_lines):
    # Use a normal paragraph with preserved newlines and mono-like formatting
    p = doc.add_paragraph()
    for i, line in enumerate(code_lines.split('\n')):
        run = p.add_run(("" if i == 0 else "\n") + line)
        # Keep Times New Roman to stay within spec; use same size
        set_run_font(run, size_pt=12)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_figure_placeholder(doc, number, caption):
    # Placeholder text; user can replace with actual screenshot
    p = doc.add_paragraph()
    r = p.add_run(f"Figure {number}: {caption}\n[Insert screenshot here: docs/screenshots/figure_{number}.png]")
    set_run_font(r, size_pt=12)
    p.paragraph_format.line_spacing = 1.5


def build_document(output_path, front_page=None):
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()

    # Default body style font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(12)

    # Front Page
    fp = front_page or {
        'title': 'Grade Manager – Cloud Computing Mini Project',
        'name': 'Name',
        'register_number': 'Register Number',
        'department': 'Department',
        'year_semester': 'Year / Semester',
        'batch': 'Batch',
        'subject_name': 'Subject Name',
        'subject_code': 'Subject Code',
        'date_of_submission': datetime.now().strftime('%Y-%m-%d'),
    }

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(fp['title'])
    set_run_font(title_run, size_pt=14, bold=True)

    for key in ['name','register_number','department','year_semester','batch','subject_name','subject_code','date_of_submission']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{key.replace('_',' ').title()}: {fp[key]}")
        set_run_font(r, size_pt=12)
        p.paragraph_format.line_spacing = 1.5

    doc.add_page_break()

    # Title
    add_heading(doc, 'Grade Manager – Cloud Computing Mini Project', level=1)

    # Abstract
    add_heading(doc, 'Abstract', level=1)
    add_body_paragraph(doc, (
        'Aim: Build a grade management system to calculate attendance percentage, '
        'compute combined internal/external marks, assign grades, and generate reports.\n'
        'Technology: React (frontend), Python FastAPI/Flask-style services with SQLAlchemy (backend), '
        'MySQL database, optional AWS S3 for storage.\n'
        'Tools: SQLAlchemy, python-docx for documentation, HTML report generator.\n'
        'Outcome: End-to-end evaluation pipeline with report generation and a searchable UI.'
    ))

    # Introduction
    add_heading(doc, 'Introduction', level=1)
    add_body_paragraph(doc, (
        'Tools Used: React for UI, a Python web framework with SQLAlchemy ORM for data access, '
        'MySQL for persistent storage, and HTML/CSS for report generation.\n'
        'Objective: Provide a reliable and configurable grade evaluation workflow.\n'
        'Scope: Subject management, criteria configuration, batch-wise evaluation, and report export.'
    ))

    # Project Workflow
    add_heading(doc, 'Project Workflow and Main Logic', level=1)
    add_body_paragraph(doc, (
        '1) Data flow: Students and subjects are loaded; criteria define internal/external weights.\n'
        '2) Attendance% is computed from Attendance records versus expected periods.\n'
        '3) Combined marks are a weighted sum of internal and external totals.\n'
        '4) Grade rules: If attendance < 75% ⇒ SA, else thresholds O/A+/A/B+/B/U.\n'
        '5) Reports are generated as styled HTML tables per subject and batch.'
    ))

    # Implementation - Backend
    add_heading(doc, 'Implementation', level=1)
    add_heading(doc, 'Backend Key Functions', level=1)

    add_heading(doc, 'Attendance Percentage', level=1)
    add_body_paragraph(doc, 'Source: backend/services/grade_calculator.py')
    add_code_block(doc, (
        'def calculate_attendance_percent(db: Session, student_id: int, subject_id: int) -> float:\n'
        '    try:\n'
        '        total_hours = db.query(func.sum(Attendance.hours)).\\\n'
        '            filter(Attendance.student_id == student_id, Attendance.subject_id == subject_id).scalar() or 0.0\n'
        '        from models import AttendancePeriod\n'
        '        total_periods = db.query(AttendancePeriod).\\\n'
        '            filter(AttendancePeriod.subject_id == subject_id).count()\n'
        '        if total_periods == 0:\n'
        '            total_records = db.query(func.count(Attendance.id)).\\\n'
        '                filter(Attendance.student_id == student_id, Attendance.subject_id == subject_id).scalar() or 0\n'
        '            return 100.0 if total_records > 0 else 0.0\n'
        '        total_expected_hours = float(total_periods)\n'
        '        if total_expected_hours == 0:\n'
        '            return 0.0\n'
        '        attendance_percent = (float(total_hours) / total_expected_hours) * 100\n'
        '        return min(attendance_percent, 100.0)\n'
        '    except Exception:\n'
        '        return 0.0'
    ))

    add_heading(doc, 'Combined Marks', level=1)
    add_body_paragraph(doc, 'Source: backend/services/grade_calculator.py')
    add_code_block(doc, (
        'def calculate_combined_marks(internal_total, external_total, internal_weight):\n'
        '    external_weight = 100 - internal_weight\n'
        '    combined = (internal_total * (internal_weight / 100.0)) + \\\n'
        '               (external_total * (external_weight / 100.0))\n'
        '    return round(combined, 2)'
    ))

    add_heading(doc, 'Grade Rules', level=1)
    add_body_paragraph(doc, 'Source: backend/services/grade_calculator.py')
    add_code_block(doc, (
        'def calculate_grade(combined_marks, attendance_percent):\n'
        '    if attendance_percent < 75.0:\n'
        '        return "SA"\n'
        '    if combined_marks > 90:\n'
        '        return "O"\n'
        '    elif combined_marks > 80:\n'
        '        return "A+"\n'
        '    elif combined_marks > 70:\n'
        '        return "A"\n'
        '    elif combined_marks > 60:\n'
        '        return "B+"\n'
        '    elif combined_marks > 50:\n'
        '        return "B"\n'
        '    else:\n'
        '        return "U"'
    ))

    add_heading(doc, 'HTML Report Generation (excerpt)', level=1)
    add_body_paragraph(doc, 'Source: backend/services/report_generator.py')
    add_code_block(doc, (
        '<tr>\n'
        '  <td>{eval.student_ref.regno}</td>\n'
        '  <td>{eval.student_ref.name}</td>\n'
        '  <td style="text-align:center;">{eval.attendance_percent:.2f}%</td>\n'
        '  <td style="text-align:center;">{eval.internal_total:.2f}</td>\n'
        '  <td style="text-align:center;">{eval.external_total:.2f}</td>\n'
        '  <td style="text-align:center;">{eval.combined_marks:.2f}</td>\n'
        '  <td style="text-align:center; color:{grade_color}; font-weight:bold;">{eval.grade}</td>\n'
        '</tr>'
    ))

    # Implementation - Frontend
    add_heading(doc, 'Frontend Components', level=1)
    add_body_paragraph(doc, (
        'Key components include GradeEvaluationTable (subject selection, batch, criteria, evaluate & list results),\n'
        'CriteriaManager (configure internal components), EvaluateButton (trigger evaluation),\n'
        'ReportsList (list and open generated reports).'
    ))

    # Screenshots placeholders
    add_heading(doc, 'Implementation Screenshots', level=1)
    add_figure_placeholder(doc, 1, 'Frontend Home (App)')
    add_figure_placeholder(doc, 2, 'Grade Evaluation Table populated')
    add_figure_placeholder(doc, 3, 'Criteria Manager')
    add_figure_placeholder(doc, 4, 'Evaluate Button action')
    add_figure_placeholder(doc, 5, 'Reports List')
    add_figure_placeholder(doc, 6, 'Sample HTML Report Preview')

    # Results
    add_heading(doc, 'Results', level=1)
    add_body_paragraph(doc, 'Provide screenshots demonstrating successful evaluation and report generation.')

    # Conclusion
    add_heading(doc, 'Conclusion', level=1)
    add_body_paragraph(doc, 'The Grade Manager streamlines evaluation workflows with configurable criteria and automated report generation.')

    # References
    add_heading(doc, 'References', level=1)
    add_body_paragraph(doc, (
        '1) React Documentation – https://react.dev\n'
        '2) SQLAlchemy – https://docs.sqlalchemy.org\n'
        '3) FastAPI – https://fastapi.tiangolo.com / Flask – https://flask.palletsprojects.com\n'
        '4) python-docx – https://python-docx.readthedocs.io\n'
        '5) MySQL – https://dev.mysql.com/doc/'
    ))

    # Tools/Experiments format
    add_heading(doc, 'Tools / Experiments', level=1)

    # Example tool entry
    add_heading(doc, '1) Tool: SQLAlchemy ORM', level=1)
    add_body_paragraph(doc, 'Introduction: SQLAlchemy provides ORM and Core for Python data access.')
    add_body_paragraph(doc, 'Purpose: Model database entities, manage sessions, and perform queries efficiently.')
    add_body_paragraph(doc, 'Deployment of a single feature: Configure SessionLocal and engine in backend/config.py.')
    add_figure_placeholder(doc, 7, 'Code snippet or screenshot of configuration')
    add_body_paragraph(doc, 'Conclusion: Simplifies database operations and improves maintainability.')

    # Another example tool entry
    add_heading(doc, '2) Tool: HTML Report Generator', level=1)
    add_body_paragraph(doc, 'Introduction: Server-side HTML generation for printable evaluation reports.')
    add_body_paragraph(doc, 'Purpose: Produce consistent, shareable reports per subject and batch.')
    add_body_paragraph(doc, 'Deployment of a single feature: Generate table with grades and totals.')
    add_figure_placeholder(doc, 8, 'Rendered report screenshot')
    add_body_paragraph(doc, 'Conclusion: Provides a ready-to-print artifact for evaluations.')

    # Formatting note
    note = doc.add_paragraph()
    nr = note.add_run('\nNote: Use black & white printing, staple neatly. Ensure titles 14pt, body 12pt, 1.5 spacing.')
    set_run_font(nr, size_pt=12)

    doc.save(output_path)


if __name__ == '__main__':
    out = os.path.join(os.path.dirname(__file__), '..', 'docs', 'GradeManager_MiniProject_Documentation.docx')
    out = os.path.abspath(out)
    build_document(out)
    print(f"Generated: {out}")
